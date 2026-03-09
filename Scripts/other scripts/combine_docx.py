#!/usr/bin/env python3
"""
Merge all DOCX files from the current working directory into a single DOCX file.
"""

from pathlib import Path
from docx import Document

OUTPUT_BASENAME = "COMBINED_NOTES.docx"

def main():
    root_dir = Path.cwd()
    output_file = root_dir / OUTPUT_BASENAME
    # Recursively include DOCX files from cwd and all subfolders.
    docx_files = sorted(
        p for p in root_dir.rglob("*.docx")
        if p != output_file
    )

    if not docx_files:
        print(f"No DOCX files found in: {root_dir}")
        return

    print(f"Found {len(docx_files)} DOCX files in: {root_dir}")
    combined_doc = Document()

    for i, docx_path in enumerate(docx_files, 1):
        rel_path = docx_path.relative_to(root_dir)
        print(f"[{i}/{len(docx_files)}] Processing: {rel_path}")

        combined_doc.add_heading(f"FILE: {rel_path}", level=2)
        try:
            source_doc = Document(docx_path)
            for paragraph in source_doc.paragraphs:
                combined_doc.add_paragraph(paragraph.text)
        except Exception as e:
            combined_doc.add_paragraph(f"[ERROR reading {docx_path.name}: {e}]")

        combined_doc.add_paragraph("")
        combined_doc.add_paragraph("")

    combined_doc.save(output_file)
    print(f"\nDone. Combined DOCX saved to: {output_file}")

if __name__ == "__main__":
    main()
